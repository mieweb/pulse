#!/usr/bin/env python3
"""Generate Information_Sheet_Pulse_Vault_User_Study.docx for IRB Local Site Documents."""
from pathlib import Path
try:
    from docx import Document
except ImportError:
    print("Run: pip install python-docx")
    raise

def main():
    doc = Document()
    doc.add_heading("RESEARCH PARTICIPANT Information Sheet", level=0)
    doc.add_paragraph(
        "Adoption and Effectiveness of Pulse Vault for Institutional Knowledge Sharing: "
        "A Longitudinal User Study with Employees at Participating Organizations"
    )
    doc.add_paragraph("Adolfo Coronado")
    doc.add_paragraph("Department of Computer Science")
    doc.add_paragraph("STUDY2026-00000327")
    doc.add_paragraph("Purdue University")
    doc.add_paragraph()

    doc.add_paragraph(
        "●      You are being asked to participate or be a part of a research study. "
        "Your participation is voluntary which means that you may choose not to participate at any time."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "●      The researchers hope to learn more about how employees at participating organizations use Pulse Vault — "
        "a video-based knowledge-sharing platform — and how it compares to text-based documentation. "
        "The study team will approach different organizations to implement this study. The purpose is to understand "
        "adoption, usability, and whether video documentation is helpful in your workflow."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "●      You are being asked to participate because the study team will approach different organizations to participate. "
        "If your organization agrees to participate, you are being asked because you are an employee who may use Pulse Vault "
        "for work-related knowledge sharing. Recruitment within each organization is internal (e.g., in coordination with "
        "organizational leadership); participation is voluntary."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "●      You will be asked to participate in naturalistic use of Pulse Vault (upload, view, share short videos for work), "
        "complete up to three short online surveys (baseline, midpoint, end of study) about your role, Pulse Vault usability and "
        "satisfaction, and video vs. text documentation, and optionally participate in a one-on-one interview (about 30–60 minutes) "
        "to discuss your experience."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "●      The study will take a total of: surveys about 5–15 minutes each (three surveys total); optional interview 30–60 minutes; "
        "naturalistic use for at least one month with no minimum usage required."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "Please take time to review the rest of the information. This will give you information about this study to help you decide if you want to participate."
    )
    doc.add_paragraph(
        "This study is only intended for people who are 18 years of age or older.  Do not complete the study if you are not legally considered an adult."
    )
    doc.add_paragraph(
        "Before agreeing to participate, consider the risks and potential benefits of taking part in this study."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "All research carries the risk of breach of confidentiality which means that someone outside of our study could figure out that you were in the study or information was yours.  How we will protect your information to reduce this risk is below."
    )
    doc.add_paragraph(
        "Some questions could make you feel uncomfortable. You can skip any of these questions or stop answering."
    )
    doc.add_paragraph(
        "It is unlikely that there will be personal benefits to you for participating. Having more information from the answers in this research study might help us or other researchers understand more about video-based knowledge sharing and adoption of Pulse Vault in organizational settings."
    )
    doc.add_paragraph()

    doc.add_heading("Will I receive payment or other incentive?", level=1)
    doc.add_paragraph("You will not be paid for being in this research study.")
    doc.add_paragraph()

    doc.add_heading("How will the researchers protect my information, privacy, and confidentiality?", level=1)
    doc.add_paragraph(
        "Your personal information may be shared outside the research study if required by law. We also may need to share your research records with other groups for quality assurance or data analysis. These groups include the Purdue University Institutional Review Board or its designees, and state or federal agencies who may need to access the research records (as allowed by law)."
    )
    doc.add_paragraph(
        "Your research information may also be shared with collaborators on this research at Purdue University Fort Wayne. There is no study sponsor; this is thesis research."
    )
    doc.add_paragraph(
        "The study team plans to keep answers for this study to answer research questions.  We will keep this information until we are done with the study, approximately one year after the study ends (and for at least three years after we are finished, per Purdue policy).  We may share the anonymous data and findings with other researchers or in research papers or presentations."
    )
    doc.add_paragraph(
        "If you take part in an optional interview, we may audio-record with your separate permission. Only the study team will have access to recordings; they will be stored securely and deleted or de-identified per IRB policy. Recordings will not be shared in publications."
    )
    doc.add_paragraph()

    doc.add_heading("What are my rights as a research participant in this study?", level=1)
    doc.add_paragraph(
        "You do not have to participate in this research project.  If you agree to participate, you may withdraw your participation at any time without penalty."
    )
    doc.add_paragraph()

    doc.add_heading("Who can I contact if I have questions about the study?", level=1)
    doc.add_paragraph(
        "If you have questions, comments or concerns about this research project, you can talk to one of the researchers.  Please contact the Principal Investigator: Adolfo Coronado — acoronad@pfw.edu. Faculty Advisor: Romael Haque — mdromael.haque@pfw.edu. Student Investigator: Priyam More — morepr01@pfw.edu. First point of contact: Principal Investigator."
    )
    doc.add_paragraph(
        "To report anonymously via Purdue's Hotline, see www.purdue.edu/hotline"
    )
    doc.add_paragraph(
        "If you have questions about your rights while taking part in the study or have concerns about the treatment of research participants, please call the Human Research Protection Program at (765) 494-5942, email (irb@purdue.edu) or write to:"
    )
    doc.add_paragraph()
    doc.add_paragraph("Human Research Protection Program")
    doc.add_paragraph("Purdue University")
    doc.add_paragraph("Seng Liang Wang Hall")
    doc.add_paragraph("516 Northwestern Ave")
    doc.add_paragraph("West Lafayette, IN 47906")
    doc.add_paragraph()

    doc.add_paragraph(
        "☐By clicking/checking this box, I agree to take part in this research. I am 18 years of age or older and understand the information above about my participation."
    )

    out = Path(__file__).parent / "Information_Sheet_Pulse_Vault_User_Study.docx"
    doc.save(str(out))
    print(f"Saved: {out}")

if __name__ == "__main__":
    main()
