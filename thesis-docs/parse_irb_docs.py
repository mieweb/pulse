#!/usr/bin/env python3
"""
Parse IRB .doc and .docx templates and extract text/structure.
Outputs plain text and (for .docx) paragraph-level structure for filling from thesis-docs.
"""

import os
import re
import subprocess
import sys
from pathlib import Path

# Optional: python-docx for .docx
try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def read_docx(path: str) -> dict:
    """Read .docx and return structured content (paragraphs, tables)."""
    if not HAS_DOCX:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    doc = Document(path)
    out = {"path": path, "format": "docx", "paragraphs": [], "tables": [], "full_text": []}
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            out["paragraphs"].append({"style": para.style.name, "text": text})
            out["full_text"].append(text)
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([(c.text or "").strip() for c in row.cells])
        out["tables"].append({"index": i, "rows": rows})
        for row in rows:
            out["full_text"].append(" | ".join(row))
    return out


def read_doc(path: str) -> dict:
    """Read legacy .doc using textutil (macOS) or antiword if available."""
    path = os.path.abspath(path)
    if not os.path.isfile(path):
        return {"error": f"File not found: {path}"}
    # Try textutil on macOS first
    try:
        r = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", path],
            capture_output=True,
            text=True,
            timeout=10,
        )
        if r.returncode == 0 and r.stdout:
            lines = [s.strip() for s in r.stdout.splitlines() if s.strip()]
            return {
                "path": path,
                "format": "doc",
                "paragraphs": [{"text": line} for line in lines],
                "tables": [],
                "full_text": lines,
            }
    except FileNotFoundError:
        pass
    # Fallback: try antiword (Linux/some installs)
    try:
        r = subprocess.run(
            ["antiword", path],
            capture_output=True,
            text=True,
            timeout=10,
        )
        if r.returncode == 0 and r.stdout:
            lines = [s.strip() for s in r.stdout.splitlines() if s.strip()]
            return {
                "path": path,
                "format": "doc",
                "paragraphs": [{"text": line} for line in lines],
                "tables": [],
                "full_text": lines,
            }
    except FileNotFoundError:
        pass
    return {
        "error": "Could not read .doc: install textutil (macOS) or antiword (e.g. brew install antiword)",
        "path": path,
    }


def parse_irb_file(path: str) -> dict:
    """Dispatch to .doc or .docx reader by extension."""
    path = path.strip()
    if not os.path.isabs(path):
        # Resolve relative to script dir, then cwd
        base = Path(__file__).parent.parent
        for candidate in (base / path, Path(path), Path.cwd() / path):
            if candidate.exists():
                path = str(candidate)
                break
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx(path)
    if ext == ".doc":
        return read_doc(path)
    return {"error": f"Unsupported extension: {ext}. Use .doc or .docx"}


def format_output(data: dict, want_markdown: bool = True) -> str:
    """Format parsed data as readable text/markdown."""
    if "error" in data:
        return f"Error: {data['error']}\n"
    lines = [f"# Parsed: {os.path.basename(data['path'])} ({data['format']})\n"]
    if want_markdown and data.get("paragraphs"):
        lines.append("## Paragraphs\n")
        for i, p in enumerate(data["paragraphs"]):
            style = p.get("style", "")
            text = p.get("text", "")
            if style:
                lines.append(f"- [{style}] {text}")
            else:
                lines.append(f"- {text}")
        lines.append("")
    if data.get("tables"):
        lines.append("## Tables\n")
        for t in data["tables"]:
            for row in t["rows"]:
                lines.append(" | ".join(row))
            lines.append("")
    lines.append("## Full text (concatenated)\n")
    lines.append("\n".join(data.get("full_text", [])))
    return "\n".join(lines)


def main():
    # Default: parse both IRB files from repo root
    repo_root = Path(__file__).parent.parent
    default_doc = repo_root / "HRP- 502b- Exempt Research Information Sheet.doc"
    default_docx = repo_root / "2026-01-21 HRP-503c Protocol - Exempt Research (1).docx"

    if len(sys.argv) > 1:
        paths = sys.argv[1:]
    else:
        paths = []
        if default_doc.exists():
            paths.append(str(default_doc))
        if default_docx.exists():
            paths.append(str(default_docx))
        if not paths:
            print("Usage: python parse_irb_docs.py [file1.doc] [file2.docx] ...")
            print("Or place HRP-502b .doc and HRP-503c .docx in repo root and run with no args.")
            sys.exit(1)

    for path in paths:
        print(f"\n{'='*60}\nParsing: {path}\n{'='*60}")
        data = parse_irb_file(path)
        print(format_output(data))


if __name__ == "__main__":
    main()
