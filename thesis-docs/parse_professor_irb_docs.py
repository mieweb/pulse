#!/usr/bin/env python3
"""
Parse professor's IRB-related .docx files (Consent Form, Study Protocol) and output
plain text for reference. These are from a different study (mental health app) but
provide structural/tone reference for Purdue IRB submissions.
Output is written to professor_irb_reference.txt in thesis-docs.
"""
from pathlib import Path
import sys

try:
    from docx import Document
except ImportError:
    print("Install python-docx: pip install python-docx")
    sys.exit(1)

REPO_ROOT = Path(__file__).resolve().parent.parent
PROF_DIR = REPO_ROOT / "professor" / "Evaluaion Study"  # note: typo in folder name
OUTPUT = Path(__file__).resolve().parent / "professor_irb_reference.txt"

FILES = [
    "Consent Form.docx",
    "Study Protocol_.docx",
]


def extract_docx(path: Path) -> list[str]:
    lines = []
    doc = Document(str(path))
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return lines


def main():
    out_lines = [
        "Reference text extracted from professor's IRB-related documents.",
        "These are from a different study (mental health app recommender) and are",
        "for structural/tone reference only. Pulse Vault study content is in",
        "HRP_503c_FINAL_PROTOCOL.md and HRP_502b_INFORMATION_SHEET_FINAL.md.",
        "",
        "=" * 72,
    ]
    for name in FILES:
        path = PROF_DIR / name
        if not path.exists():
            out_lines.append(f"\n[File not found: {path}]\n")
            continue
        out_lines.append(f"\n### {name}\n")
        out_lines.extend(extract_docx(path))
        out_lines.append("")
    OUTPUT.write_text("\n".join(out_lines), encoding="utf-8")
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
