#!/usr/bin/env python3
"""Generate Survey 1–3 as .docx for IRB or direct use. Requires python-docx."""
from pathlib import Path
try:
    from docx import Document
except ImportError:
    print("Run: pip install python-docx")
    raise

OUT_DIR = Path(__file__).parent / "surveys"


def add_para(doc, text):
    if text.strip():
        doc.add_paragraph(text.strip())


def baseline(doc):
    doc.add_heading("Survey 1: Baseline (after consent)", level=0)
    doc.add_paragraph("Study: Adoption and Effectiveness of Pulse Vault for Institutional Knowledge Sharing")
    doc.add_paragraph("Est. time: 3–5 minutes")
    doc.add_paragraph("When: After consent, before or right after first use of Pulse Vault")
    doc.add_paragraph()
    doc.add_heading("Instructions for participants", level=1)
    doc.add_paragraph(
        "This short survey asks about your role and your past use of video for work documentation. "
        "Your answers help us understand our participants and compare attitudes before and after the study."
    )
    doc.add_paragraph()
    doc.add_heading("Questions", level=1)
    doc.add_heading("Demographics / Role", level=2)
    doc.add_paragraph("Q1. What best describes your role at your organization?")
    doc.add_paragraph("Developer / Engineer", style="List Bullet")
    doc.add_paragraph("Support / Customer success", style="List Bullet")
    doc.add_paragraph("Product / Project management", style="List Bullet")
    doc.add_paragraph("Other: _____________", style="List Bullet")
    doc.add_paragraph("Response type: Single choice (dropdown or radio) + \"Other\" short text")
    doc.add_paragraph()
    doc.add_paragraph("Q2. Before this study, how often did you use video (e.g., screen recordings, Loom, Pulse Vault, or other) for work-related documentation or training?")
    doc.add_paragraph("Never", style="List Bullet")
    doc.add_paragraph("Rarely (once or twice)", style="List Bullet")
    doc.add_paragraph("Sometimes (a few times a month)", style="List Bullet")
    doc.add_paragraph("Often (weekly or more)", style="List Bullet")
    doc.add_paragraph("Response type: Single choice")
    doc.add_paragraph()
    doc.add_heading("Preference (for before/after comparison)", level=2)
    doc.add_paragraph("Q3. Before this study, for learning a new work procedure or tool, I preferred:")
    doc.add_paragraph("Text (wiki, docs, email)", style="List Bullet")
    doc.add_paragraph("No strong preference", style="List Bullet")
    doc.add_paragraph("Video (e.g., screen recordings, tutorials)", style="List Bullet")
    doc.add_paragraph("Response type: Single choice")
    doc.add_paragraph()
    doc.add_heading("Optional: link to usage (only if consent allows)", level=2)
    doc.add_paragraph("Q4. [Optional] If you're okay linking this survey to your Pulse Vault account for analysis, enter the email you use for Pulse Vault: _____________")
    doc.add_paragraph("(Leave blank if you prefer to stay anonymous.)")
    doc.add_paragraph("Response type: Short answer (optional); can be left blank")
    doc.add_paragraph()
    doc.add_heading("End of Survey 1", level=1)
    doc.add_paragraph("Thank you. You will receive the midpoint survey in about 2 weeks.")


def midpoint(doc):
    doc.add_heading("Survey 2: Midpoint (~2 weeks into study)", level=0)
    doc.add_paragraph("Study: Adoption and Effectiveness of Pulse Vault for Institutional Knowledge Sharing")
    doc.add_paragraph("Est. time: 5–8 minutes")
    doc.add_paragraph("When: Approximately 2 weeks after you started using Pulse Vault")
    doc.add_paragraph()
    doc.add_heading("Instructions for participants", level=1)
    doc.add_paragraph(
        "For the next 10 questions, \"this system\" means Pulse Vault (the video platform you've been using). "
        "Rate each statement from 1 = Strongly disagree to 5 = Strongly agree."
    )
    doc.add_paragraph()
    doc.add_heading("System Usability Scale (SUS)", level=1)
    doc.add_paragraph("Scale for all items below: 1 = Strongly disagree, 2, 3, 4, 5 = Strongly agree")
    doc.add_paragraph()
    sus_mid = [
        "Q5. I think that I would like to use this system frequently.",
        "Q6. I found the system unnecessarily complex.",
        "Q7. I thought the system was easy to use.",
        "Q8. I think that I would need the support of a technical person to be able to use this system.",
        "Q9. I found the various functions in this system were well integrated.",
        "Q10. I thought there was too much inconsistency in this system.",
        "Q11. I would imagine that most people would learn to use this system very quickly.",
        "Q12. I found the system very cumbersome to use.",
        "Q13. I felt very confident using the system.",
        "Q14. I needed to learn a lot of things before I could get going with this system.",
    ]
    for q in sus_mid:
        doc.add_paragraph(q)
    doc.add_paragraph("Response type for Q5–Q14: Linear scale 1–5 (Strongly disagree → Strongly agree)")
    doc.add_paragraph()
    doc.add_heading("Satisfaction", level=1)
    doc.add_paragraph("Q15. Overall, I am satisfied with Pulse Vault for sharing work knowledge through video.")
    doc.add_paragraph("1 – Strongly disagree    2    3    4    5 – Strongly agree")
    doc.add_paragraph("Response type: Linear scale 1–5")
    doc.add_paragraph()
    doc.add_paragraph("Q16. [Optional] Anything you'd like to share about your experience so far?")
    doc.add_paragraph("Response type: Long answer (paragraph); optional")
    doc.add_paragraph()
    doc.add_heading("End of Survey 2", level=1)
    doc.add_paragraph("Thank you. You will receive the final survey in about 2 more weeks (around 1 month from the start of the study).")


def end_of_study(doc):
    doc.add_heading("Survey 3: End of study (~1 month)", level=0)
    doc.add_paragraph("Study: Adoption and Effectiveness of Pulse Vault for Institutional Knowledge Sharing")
    doc.add_paragraph("Est. time: 8–12 minutes")
    doc.add_paragraph("When: Approximately 1 month after you started, or when the study closes")
    doc.add_paragraph()
    doc.add_heading("Instructions for participants", level=1)
    doc.add_paragraph(
        "This is the final survey. It includes usability and satisfaction questions again, plus questions about video vs. text documentation. "
        "\"This system\" means Pulse Vault. Rate scale: 1 = Strongly disagree to 5 = Strongly agree unless otherwise indicated."
    )
    doc.add_paragraph()
    doc.add_heading("System Usability Scale (SUS) — same 10 items as midpoint", level=1)
    doc.add_paragraph("Scale: 1 = Strongly disagree … 5 = Strongly agree")
    doc.add_paragraph()
    sus_end = [
        "Q17. I think that I would like to use this system frequently.",
        "Q18. I found the system unnecessarily complex.",
        "Q19. I thought the system was easy to use.",
        "Q20. I think that I would need the support of a technical person to be able to use this system.",
        "Q21. I found the various functions in this system were well integrated.",
        "Q22. I thought there was too much inconsistency in this system.",
        "Q23. I would imagine that most people would learn to use this system very quickly.",
        "Q24. I found the system very cumbersome to use.",
        "Q25. I felt very confident using the system.",
        "Q26. I needed to learn a lot of things before I could get going with this system.",
    ]
    for q in sus_end:
        doc.add_paragraph(q)
    doc.add_paragraph("Response type: Linear scale 1–5")
    doc.add_paragraph()
    doc.add_heading("Satisfaction", level=1)
    doc.add_paragraph("Q27. Overall, I am satisfied with Pulse Vault for sharing work knowledge through video.")
    doc.add_paragraph("1 – Strongly disagree … 5 – Strongly agree")
    doc.add_paragraph("Response type: Linear scale 1–5")
    doc.add_paragraph()
    doc.add_heading("Video vs. text (research questions)", level=1)
    doc.add_paragraph("Q28. For learning a new work procedure or tool, I prefer:")
    doc.add_paragraph("1 = Strongly prefer text (wiki, docs, email)", style="List Bullet")
    doc.add_paragraph("2 = Prefer text", style="List Bullet")
    doc.add_paragraph("3 = No preference", style="List Bullet")
    doc.add_paragraph("4 = Prefer video (e.g., Pulse Vault, screen recordings)", style="List Bullet")
    doc.add_paragraph("5 = Strongly prefer video", style="List Bullet")
    doc.add_paragraph("Response type: Single choice or scale 1–5")
    doc.add_paragraph()
    doc.add_paragraph("Q29. \"Video documentation on Pulse Vault was more helpful than text documentation when I needed to learn a new procedure or fix an issue.\"")
    doc.add_paragraph("1 – Strongly disagree … 5 – Strongly agree")
    doc.add_paragraph("Response type: Linear scale 1–5")
    doc.add_paragraph()
    doc.add_paragraph("Q30. \"I would choose to look at a short video before reading a text doc when both were available for the same topic.\"")
    doc.add_paragraph("1 – Strongly disagree … 5 – Strongly agree")
    doc.add_paragraph("Response type: Linear scale 1–5")
    doc.add_paragraph()
    doc.add_paragraph("Q31. In the past month (or study period), when you needed to learn something work-related (procedure, tool, bug fix), how often did you use:")
    doc.add_paragraph("Only text (wiki, docs, email)", style="List Bullet")
    doc.add_paragraph("Mostly text", style="List Bullet")
    doc.add_paragraph("Both about equally", style="List Bullet")
    doc.add_paragraph("Mostly video (including Pulse Vault)", style="List Bullet")
    doc.add_paragraph("Only video", style="List Bullet")
    doc.add_paragraph("Response type: Single choice")
    doc.add_paragraph()
    doc.add_paragraph("Q32. Roughly how many work procedures or tutorials did you create in the study period?")
    doc.add_paragraph("As VIDEO (e.g., uploaded to Pulse Vault): _____________")
    doc.add_paragraph("(number, or: None / 1–2 / 3–5 / 6–10 / More than 10)")
    doc.add_paragraph("As TEXT (wiki, doc, email, Slack): _____________")
    doc.add_paragraph("(same options)")
    doc.add_paragraph("Response type: Two short-answer or dropdown fields")
    doc.add_paragraph()
    doc.add_heading("Adoption barriers and enablers", level=1)
    doc.add_paragraph("Q33. What helped you use Pulse Vault? (Check all that apply or add your own.)")
    doc.add_paragraph("Easy to upload / share", style="List Bullet")
    doc.add_paragraph("Fits my workflow", style="List Bullet")
    doc.add_paragraph("Team encouraged it", style="List Bullet")
    doc.add_paragraph("Faster than writing docs", style="List Bullet")
    doc.add_paragraph("Other: _____________", style="List Bullet")
    doc.add_paragraph("Response type: Multiple choice (checkboxes) + optional \"Other\" short text")
    doc.add_paragraph()
    doc.add_paragraph("Q34. What made it hard to use Pulse Vault or to use it more? (Check all that apply or add your own.)")
    doc.add_paragraph("Too time-consuming", style="List Bullet")
    doc.add_paragraph("Didn't know when to use it", style="List Bullet")
    doc.add_paragraph("Prefer text", style="List Bullet")
    doc.add_paragraph("Technical issues", style="List Bullet")
    doc.add_paragraph("Nothing – it was fine", style="List Bullet")
    doc.add_paragraph("Other: _____________", style="List Bullet")
    doc.add_paragraph("Response type: Multiple choice (checkboxes) + optional \"Other\" short text")
    doc.add_paragraph()
    doc.add_paragraph("Q35. [Optional] Any other feedback about Pulse Vault or video vs. text documentation at work?")
    doc.add_paragraph("Response type: Long answer (paragraph); optional")
    doc.add_paragraph()
    doc.add_heading("Optional: link to usage", level=1)
    doc.add_paragraph("Q36. [Optional] If you're okay linking this survey to your Pulse Vault account for analysis, enter the email you use for Pulse Vault: _____________")
    doc.add_paragraph("(Leave blank to stay anonymous.)")
    doc.add_paragraph("Response type: Short answer; optional")
    doc.add_paragraph()
    doc.add_heading("End of Survey 3", level=1)
    doc.add_paragraph("Thank you for participating. If you agreed to an optional interview, we will contact you separately.")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = [
        ("01_Baseline_Survey.docx", baseline),
        ("02_Midpoint_Survey.docx", midpoint),
        ("03_End_of_Study_Survey.docx", end_of_study),
    ]
    for name, build_fn in files:
        doc = Document()
        build_fn(doc)
        out = OUT_DIR / name
        doc.save(str(out))
        print(f"Saved: {out}")
    print("Done. All three survey .docx files are in thesis-docs/surveys/")


if __name__ == "__main__":
    main()
