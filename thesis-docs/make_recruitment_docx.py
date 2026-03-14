#!/usr/bin/env python3
"""Generate Recruitment_Email_Pulse_Vault.docx for IRB upload."""
from pathlib import Path
try:
    from docx import Document
except ImportError:
    print("Run: pip install python-docx")
    raise

def main():
    doc = Document()
    doc.add_heading("Recruitment Email – Pulse Vault User Study", level=0)
    doc.add_paragraph()
    doc.add_paragraph("Use the text below when recruiting participants (e.g., by email to organizational contacts or employees). Do not use until IRB approval has been received.")
    doc.add_paragraph()
    doc.add_paragraph("—" * 20)
    doc.add_paragraph()
    doc.add_paragraph("Subject: Invitation to participate in Pulse Vault research study")
    doc.add_paragraph()
    doc.add_paragraph(
        "You are invited to participate in a research study about how employees use Pulse Vault—a video-based knowledge-sharing platform—for work-related documentation and knowledge sharing. The study is conducted by researchers at Purdue University Fort Wayne and is part of a Master's thesis.")
    doc.add_paragraph()
    doc.add_paragraph(
        "Participation involves: using Pulse Vault as you normally would for at least one month; completing up to three short online surveys (about 5–15 minutes each); and optionally taking part in a one-on-one interview (about 30–60 minutes) to discuss your experience.")
    doc.add_paragraph()
    doc.add_paragraph(
        "Participation is voluntary and there is no compensation. You may skip any question or withdraw at any time.")
    doc.add_paragraph()
    doc.add_paragraph(
        "If you are interested, please contact us to receive the study information sheet and to take part. Questions can be directed to:")
    doc.add_paragraph()
    doc.add_paragraph("Principal Investigator: Adolfo Coronado — acoronad@pfw.edu | 260-481-6181")
    doc.add_paragraph("Student Investigator: Priyam More — morepr01@pfw.edu | 260-449-1394")
    doc.add_paragraph()
    doc.add_heading("Where to access Pulse Vault and Pulse", level=1)
    doc.add_paragraph("To get started with the study, you can use Pulse Vault (web) and/or the Pulse app (mobile):")
    doc.add_paragraph()
    doc.add_paragraph("Pulse Vault (web): https://pulse-vault.opensource.mieweb.org/")
    doc.add_paragraph("Pulse app (iOS): https://apps.apple.com/app/pulse-video-recorder/id6748621024")
    doc.add_paragraph("Pulse app (Android): https://play.google.com/store/apps/details?id=com.mieweb.pulse")
    doc.add_paragraph("Pulse (GitHub): https://github.com/mieweb/pulse")
    doc.add_paragraph("Pulse Vault (GitHub): https://github.com/mieweb/pulsevault")
    doc.add_paragraph()
    doc.add_paragraph("Thank you for your consideration.")
    doc.add_paragraph()
    doc.add_paragraph("—" * 20)
    out = Path(__file__).parent / "Recruitment_Email_Pulse_Vault.docx"
    doc.save(str(out))
    print(f"Saved: {out}")

if __name__ == "__main__":
    main()
